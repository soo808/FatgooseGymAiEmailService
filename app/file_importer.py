"""
file_importer.py — Excel / docx 文件解析，输出统一 QA 列表
  导入后 question_zh / answer_zh 留空（UI 显示"⚠️ 未翻译"badge）。
  支持四种解析模式：
    qa_excel       — Excel：指定 Q 列 / A 列
    paragraph_excel— Excel：单列正文，每行一条（question_jp = answer_jp = 该行）
    docx_qa        — docx：Q/A 标记行（如 Q: / A:）交替出现
    docx_paragraphs— docx：Heading1/2 → question，后续段落 → answer
"""

from __future__ import annotations

import io
import re
import uuid
from typing import Optional


def _new_id() -> str:
    return str(uuid.uuid4())


# ─────────────────────────────────────────────────────────────────────────────
# Excel 解析
# ─────────────────────────────────────────────────────────────────────────────

def parse_qa_excel(file_bytes: bytes, q_col: str, a_col: str) -> list[dict]:
    """
    从 Excel 按列名提取 QA 对。
    q_col / a_col 为列名（如 "質問" / "回答"）。
    跳过 Q 或 A 为空的行。
    """
    import pandas as pd

    df = pd.read_excel(io.BytesIO(file_bytes), dtype=str)
    df.fillna("", inplace=True)

    if q_col not in df.columns or a_col not in df.columns:
        raise ValueError(
            f"列名不存在。文件包含: {list(df.columns)}，需要: [{q_col}, {a_col}]"
        )

    result = []
    for _, row in df.iterrows():
        q = str(row[q_col]).strip()
        a = str(row[a_col]).strip()
        if not q or not a:
            continue
        result.append({
            "id": _new_id(),
            "question_jp": q,
            "answer_jp": a,
            "question_zh": "",
            "answer_zh": "",
            "inquiry_type": "",
            "environment": "",
            "app_version": "",
            "date": "",
            "total_turns": 1,
        })
    return result


def parse_paragraph_excel(file_bytes: bytes, text_col: str) -> list[dict]:
    """
    从 Excel 单列提取段落，每行作为独立知识条目（Q = A = 该行文本）。
    适用于规则/政策类文档（无明确 Q/A 对）。
    """
    import pandas as pd

    df = pd.read_excel(io.BytesIO(file_bytes), dtype=str)
    df.fillna("", inplace=True)

    if text_col not in df.columns:
        raise ValueError(
            f"列名 [{text_col}] 不存在。文件包含: {list(df.columns)}"
        )

    result = []
    for _, row in df.iterrows():
        text = str(row[text_col]).strip()
        if not text:
            continue
        result.append({
            "id": _new_id(),
            "question_jp": text,
            "answer_jp": text,
            "question_zh": "",
            "answer_zh": "",
            "inquiry_type": "",
            "environment": "",
            "app_version": "",
            "date": "",
            "total_turns": 1,
        })
    return result


# ─────────────────────────────────────────────────────────────────────────────
# docx 解析
# ─────────────────────────────────────────────────────────────────────────────

def parse_docx_paragraphs(file_bytes: bytes) -> list[dict]:
    """
    docx 段落切分模式：
    - Heading1/Heading2 → question_jp（开始新条目）
    - 后续 Normal/Body 段落 → 追加到当前条目的 answer_jp
    - 没有 Heading 时每段独立为一条（Q = A = 该段）
    """
    from docx import Document

    doc = Document(io.BytesIO(file_bytes))
    result: list[dict] = []
    current: Optional[dict] = None
    has_heading = any(p.style.name.startswith("Heading") for p in doc.paragraphs)

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        if has_heading and para.style.name.startswith("Heading"):
            if current:
                result.append(current)
            current = {
                "id": _new_id(),
                "question_jp": text,
                "answer_jp": "",
                "question_zh": "",
                "answer_zh": "",
                "inquiry_type": "",
                "environment": "",
                "app_version": "",
                "date": "",
                "total_turns": 1,
            }
        elif has_heading and current is not None:
            current["answer_jp"] = (current["answer_jp"] + "\n" + text).strip()
        else:
            # 无 Heading：每段独立
            result.append({
                "id": _new_id(),
                "question_jp": text,
                "answer_jp": text,
                "question_zh": "",
                "answer_zh": "",
                "inquiry_type": "",
                "environment": "",
                "app_version": "",
                "date": "",
                "total_turns": 1,
            })

    if has_heading and current:
        result.append(current)

    return result


def parse_docx_qa(file_bytes: bytes, q_marker: str = "Q:", a_marker: str = "A:") -> list[dict]:
    """
    docx Q/A 标记模式：
    - 以 q_marker 开头的段落 → question_jp（去掉标记前缀）
    - 以 a_marker 开头的段落 → answer_jp（追加到当前 Q）
    - 遇到新 Q 标记时，前一组 QA 入库
    """
    from docx import Document

    doc = Document(io.BytesIO(file_bytes))
    result: list[dict] = []
    current: Optional[dict] = None

    def strip_marker(text: str, marker: str) -> str:
        return text[len(marker):].strip() if text.startswith(marker) else text.strip()

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        if text.startswith(q_marker):
            if current and current["question_jp"]:
                result.append(current)
            current = {
                "id": _new_id(),
                "question_jp": strip_marker(text, q_marker),
                "answer_jp": "",
                "question_zh": "",
                "answer_zh": "",
                "inquiry_type": "",
                "environment": "",
                "app_version": "",
                "date": "",
                "total_turns": 1,
            }
        elif text.startswith(a_marker) and current is not None:
            ans = strip_marker(text, a_marker)
            current["answer_jp"] = (current["answer_jp"] + "\n" + ans).strip()

    if current and current["question_jp"]:
        result.append(current)

    return result
